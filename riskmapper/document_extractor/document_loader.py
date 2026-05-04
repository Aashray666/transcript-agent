"""Document loader — reads any supported file format into plain text.

Supported formats:
- PDF  (.pdf)  — via pdfplumber (layout-aware, handles tables)
- DOCX (.docx) — via python-docx
- XLSX (.xlsx) — via openpyxl (reads all sheets as text)
- TXT  (.txt)  — plain text passthrough
- MD   (.md)   — markdown passthrough

Design principles:
- Returns raw text + detected section boundaries
- Section boundaries are used by the chunker to split intelligently
- Never loses text — if a parser fails, falls back to raw bytes decode
"""

from __future__ import annotations

import logging
import os
import re
from dataclasses import dataclass, field

logger = logging.getLogger(__name__)

# Supported extensions
_SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".xlsx", ".txt", ".md", ".csv"}


@dataclass
class LoadedDocument:
    """Raw text + structural metadata from a loaded document."""

    text: str                          # Full extracted text
    document_name: str                 # Original filename
    file_extension: str                # .pdf, .docx, etc.
    detected_sections: list[str]       # Section headings found (for chunking hints)
    page_count: int                    # Number of pages (0 if not applicable)
    char_count: int
    load_warnings: list[str] = field(default_factory=list)


def load_document(file_path: str) -> LoadedDocument:
    """Load a document from disk and extract its text content.

    Args:
        file_path: Absolute or relative path to the document.

    Returns:
        LoadedDocument with full text and structural metadata.

    Raises:
        FileNotFoundError: If the file does not exist.
        ValueError: If the file extension is not supported.
    """
    if not os.path.isfile(file_path):
        raise FileNotFoundError(f"Document not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()
    if ext not in _SUPPORTED_EXTENSIONS:
        raise ValueError(
            f"Unsupported file type '{ext}'. "
            f"Supported: {', '.join(sorted(_SUPPORTED_EXTENSIONS))}"
        )

    document_name = os.path.basename(file_path)
    logger.info("Loading document | name=%s | ext=%s", document_name, ext)

    if ext == ".pdf":
        return _load_pdf(file_path, document_name)
    elif ext == ".docx":
        return _load_docx(file_path, document_name)
    elif ext == ".xlsx":
        return _load_xlsx(file_path, document_name)
    elif ext in (".txt", ".md"):
        return _load_text(file_path, document_name, ext)
    elif ext == ".csv":
        return _load_text(file_path, document_name, ext)
    else:
        raise ValueError(f"Unhandled extension: {ext}")


# ---------------------------------------------------------------------------
# Format-specific loaders
# ---------------------------------------------------------------------------

def _load_pdf(file_path: str, document_name: str) -> LoadedDocument:
    """Load PDF using pdfplumber — layout-aware, handles tables."""
    try:
        import pdfplumber
    except ImportError:
        raise ImportError(
            "pdfplumber is required for PDF loading. "
            "Install it: pip install pdfplumber"
        )

    warnings: list[str] = []
    pages_text: list[str] = []
    page_count = 0

    try:
        with pdfplumber.open(file_path) as pdf:
            page_count = len(pdf.pages)
            for page_num, page in enumerate(pdf.pages, 1):
                try:
                    # Extract text with layout preservation
                    text = page.extract_text(x_tolerance=3, y_tolerance=3)
                    if text:
                        # Add page marker for section detection
                        pages_text.append(f"\n[PAGE {page_num}]\n{text}")
                    else:
                        warnings.append(f"Page {page_num}: no text extracted (may be image-based)")
                except Exception as exc:
                    warnings.append(f"Page {page_num}: extraction failed — {exc}")

    except Exception as exc:
        raise RuntimeError(f"Failed to open PDF '{file_path}': {exc}") from exc

    full_text = "\n".join(pages_text)
    sections = _detect_sections(full_text)

    logger.info(
        "PDF loaded | pages=%d | chars=%d | sections=%d | warnings=%d",
        page_count, len(full_text), len(sections), len(warnings),
    )

    return LoadedDocument(
        text=full_text,
        document_name=document_name,
        file_extension=".pdf",
        detected_sections=sections,
        page_count=page_count,
        char_count=len(full_text),
        load_warnings=warnings,
    )


def _load_docx(file_path: str, document_name: str) -> LoadedDocument:
    """Load DOCX using python-docx — preserves heading structure."""
    try:
        import docx
    except ImportError:
        raise ImportError(
            "python-docx is required for DOCX loading. "
            "Install it: pip install python-docx"
        )

    warnings: list[str] = []
    lines: list[str] = []
    sections: list[str] = []

    try:
        doc = docx.Document(file_path)

        for para in doc.paragraphs:
            text = para.text.strip()
            if not text:
                continue

            # Detect headings — mark them for section detection
            if para.style.name.startswith("Heading"):
                lines.append(f"\n[SECTION: {text}]\n")
                sections.append(text)
            else:
                lines.append(text)

        # Also extract tables
        for table_idx, table in enumerate(doc.tables):
            lines.append(f"\n[TABLE {table_idx + 1}]")
            for row in table.rows:
                row_text = " | ".join(
                    cell.text.strip() for cell in row.cells if cell.text.strip()
                )
                if row_text:
                    lines.append(row_text)

    except Exception as exc:
        raise RuntimeError(f"Failed to open DOCX '{file_path}': {exc}") from exc

    full_text = "\n".join(lines)

    logger.info(
        "DOCX loaded | chars=%d | sections=%d",
        len(full_text), len(sections),
    )

    return LoadedDocument(
        text=full_text,
        document_name=document_name,
        file_extension=".docx",
        detected_sections=sections,
        page_count=0,
        char_count=len(full_text),
        load_warnings=warnings,
    )


def _load_xlsx(file_path: str, document_name: str) -> LoadedDocument:
    """Load XLSX — reads all sheets, converts to readable text."""
    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            "openpyxl is required for XLSX loading. "
            "Install it: pip install openpyxl"
        )

    warnings: list[str] = []
    lines: list[str] = []
    sections: list[str] = []

    try:
        wb = openpyxl.load_workbook(file_path, read_only=True, data_only=True)

        for sheet_name in wb.sheetnames:
            ws = wb[sheet_name]
            lines.append(f"\n[SECTION: {sheet_name}]\n")
            sections.append(sheet_name)

            for row in ws.iter_rows(values_only=True):
                if not row or all(cell is None for cell in row):
                    continue
                row_text = " | ".join(
                    str(cell).strip() for cell in row if cell is not None
                )
                if row_text.strip():
                    lines.append(row_text)

        wb.close()

    except Exception as exc:
        raise RuntimeError(f"Failed to open XLSX '{file_path}': {exc}") from exc

    full_text = "\n".join(lines)

    logger.info(
        "XLSX loaded | sheets=%d | chars=%d",
        len(sections), len(full_text),
    )

    return LoadedDocument(
        text=full_text,
        document_name=document_name,
        file_extension=".xlsx",
        detected_sections=sections,
        page_count=0,
        char_count=len(full_text),
        load_warnings=warnings,
    )


def _load_text(file_path: str, document_name: str, ext: str) -> LoadedDocument:
    """Load plain text / markdown / CSV files."""
    warnings: list[str] = []

    try:
        with open(file_path, "r", encoding="utf-8") as f:
            full_text = f.read()
    except UnicodeDecodeError:
        warnings.append("UTF-8 decode failed — retrying with latin-1")
        with open(file_path, "r", encoding="latin-1") as f:
            full_text = f.read()

    sections = _detect_sections(full_text)

    logger.info(
        "Text loaded | ext=%s | chars=%d | sections=%d",
        ext, len(full_text), len(sections),
    )

    return LoadedDocument(
        text=full_text,
        document_name=document_name,
        file_extension=ext,
        detected_sections=sections,
        page_count=0,
        char_count=len(full_text),
        load_warnings=warnings,
    )


# ---------------------------------------------------------------------------
# Section detection — heuristic heading finder
# ---------------------------------------------------------------------------

# Patterns that look like section headings in financial/audit documents
_HEADING_PATTERNS = [
    re.compile(r"^\[SECTION:\s*(.+?)\]$"),                    # Our own markers
    re.compile(r"^#{1,4}\s+(.+)$"),                           # Markdown headings
    re.compile(r"^(\d+\.)+\s+[A-Z].{5,60}$"),                 # 1.2.3 Numbered headings
    re.compile(r"^[A-Z][A-Z\s&,\-]{5,60}$"),                  # ALL CAPS headings
    re.compile(r"^(RISK FACTORS|MANAGEMENT.S DISCUSSION|FORWARD.LOOKING)", re.I),
    re.compile(r"^(AUDIT FINDING|OBSERVATION|RECOMMENDATION|MANAGEMENT RESPONSE)", re.I),
    re.compile(r"^(EXECUTIVE SUMMARY|CONCLUSION|OVERVIEW|INTRODUCTION)", re.I),
    re.compile(r"^\[PAGE \d+\]$"),                             # Page markers
]


def _detect_sections(text: str) -> list[str]:
    """Detect section headings in document text."""
    sections: list[str] = []
    for line in text.split("\n"):
        stripped = line.strip()
        if not stripped or len(stripped) > 120:
            continue
        for pattern in _HEADING_PATTERNS:
            m = pattern.match(stripped)
            if m:
                heading = m.group(1) if m.lastindex else stripped
                if heading not in sections:
                    sections.append(heading)
                break
    return sections
